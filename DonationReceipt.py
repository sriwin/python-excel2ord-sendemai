import os
import docx
import pandas as pd
import configparser
from docx import Document
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders

from SendEmail import SendMail

def main() :
    cwd = os.getcwd()
    
    ## 
    config = configparser.ConfigParser()
    config.read("config.ini")
    ##print(config)
    excelFileName = config.get('excel','excelInputFileName')
    wordTemplateFileName = config.get('word','wordTemplateFileNme')
    ##print(excelFileName, wordTemplateFileName)

    ### Read Excel file    
    cols = [1, 2, 3]
    excelDF = pd.read_excel(excelFileName)
    excelDF = excelDF.dropna()
    ##print(excelDF.head)

    ### Read word file
    template_document = Document(wordTemplateFileName)
    doc = docx.Document()

    ### iterate excel rows
    for index, row in excelDF.iterrows():
        donarName = row['Name']
        donationAmount = row['Amount']
        donarEmailId = row['EmailID']
        donarReceiptFilePath = 'ShriShirdiSaiBabaMandirDonationReceipt.docx'
        print(index, row['Name'], row['Amount'], row['EmailID'])

        replaceDataDictionary  = {
            "{DONAR_NAME}": donarName,
            "{DONATED_AMOUNT}": "$"+str(donationAmount)
        }
        ##print(replaceDataDictionary)

        for replaceDataKey, replaceDataValue in replaceDataDictionary.items():
            for paragraph in template_document.paragraphs:
                replaceTextInWordDocument(paragraph, replaceDataKey, str(replaceDataValue))
                replaceTextInWordDocument(paragraph, replaceDataKey, str(replaceDataValue))
                template_document.save(donarReceiptFilePath)

        ### send email
        sendEmail = SendMail.fromConfig(config.get('gmail','gmailSenderId'), 
        config.get('gmail','gmailSenderPassword'), 
        config.get('gmail', 'gmailEmailSubject'))
        sendEmail.send(donarName, donarEmailId, donarReceiptFilePath)

        ### delete word file
        if os.path.exists(donarReceiptFilePath):
            os.remove(donarReceiptFilePath)

def replaceTextInWordDocument(paragraph, key, value):
    if key in paragraph.text:
        inline = paragraph.runs
        for item in inline:
            if key in item.text:
                item.text = item.text.replace(key, value)


if __name__ == '__main__':
    main()